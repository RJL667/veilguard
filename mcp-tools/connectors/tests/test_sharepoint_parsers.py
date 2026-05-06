"""Tests for SharePoint file content extraction (now LlamaIndex-backed).

The actual parsing logic lives in ``_base/parsing.py`` — this file
only verifies that the SharePoint shim re-exports the shared
``parse_to_text`` and that the LlamaIndex pipeline returns sensible
results for the formats SharePoint corpora actually contain.

LlamaIndex's individual readers are optional installs (python-docx,
pypdf, python-pptx, openpyxl, ...). Tests for a given format are
skipped when the corresponding reader isn't available — the
fallback is the placeholder, which is itself tested separately.
"""
from __future__ import annotations

import io
import pathlib
import sys

import pytest

_CONNECTORS_DIR = pathlib.Path(__file__).resolve().parent.parent
sys.path.insert(0, str(_CONNECTORS_DIR))

from sharepoint.parsers import parse_to_text  # noqa: E402


# ─── plain text family ────────────────────────────────────────────────


class TestPlainText:
    def test_txt_decoded_directly(self):
        out = parse_to_text(b"hello world\nsecond line", "notes.txt")
        # Plain-text family bypasses LlamaIndex — preserves original
        # whitespace exactly.
        assert "hello world" in out
        assert "second line" in out

    def test_json(self):
        out = parse_to_text(b'{"k": "v"}', "data.json")
        assert '"k"' in out

    def test_yaml(self):
        out = parse_to_text(b"key: value", "config.yaml")
        assert "key: value" in out

    def test_xml(self):
        out = parse_to_text(b"<root><leaf>text</leaf></root>", "doc.xml")
        assert "<leaf>text</leaf>" in out

    def test_mime_type_text_dispatches_to_decode(self):
        out = parse_to_text(b"hello", "noext", mime_type="text/plain")
        assert out == "hello"

    def test_utf8_then_latin1_fallback(self):
        # Bytes that are not valid UTF-8 — should decode via latin-1
        raw = b"\xe9clair"
        out = parse_to_text(raw, "snack.txt")
        assert "clair" in out


# ─── docx (via LlamaIndex DocxReader) ─────────────────────────────────


class TestDocx:
    def test_real_docx(self):
        try:
            import docx  # python-docx
        except ImportError:
            pytest.skip("python-docx not installed (LlamaIndex DocxReader dep)")
        try:
            from llama_index.core import SimpleDirectoryReader  # noqa: F401
        except ImportError:
            pytest.skip("llama-index-core not installed")

        buf = io.BytesIO()
        doc = docx.Document()
        doc.add_paragraph("First paragraph in the doc.")
        doc.add_paragraph("Second paragraph with searchable text.")
        doc.save(buf)
        raw = buf.getvalue()

        out = parse_to_text(raw, "doc.docx")
        # LlamaIndex DocxReader concatenates paragraphs — exact format
        # may vary by version, but content is present.
        assert "First paragraph" in out
        assert "Second paragraph" in out
        assert "searchable text" in out

    def test_corrupt_docx_returns_placeholder(self):
        try:
            from llama_index.core import SimpleDirectoryReader  # noqa: F401
        except ImportError:
            pytest.skip("llama-index-core not installed")
        out = parse_to_text(b"not a docx", "broken.docx")
        # Either a placeholder OR an empty string (LlamaIndex sometimes
        # returns empty rather than raising on malformed binary). Both
        # are acceptable graceful-degradation outcomes.
        assert out == "" or out.startswith("[binary content")


# ─── pdf (via LlamaIndex PDFReader) ───────────────────────────────────


class TestPdf:
    def test_corrupt_pdf_handled_gracefully(self):
        try:
            from llama_index.core import SimpleDirectoryReader  # noqa: F401
        except ImportError:
            pytest.skip("llama-index-core not installed")
        out = parse_to_text(b"not a pdf", "broken.pdf")
        assert out == "" or out.startswith("[binary content")


# ─── unknown / placeholder ────────────────────────────────────────────


class TestPlaceholder:
    def test_unknown_extension(self):
        # An extension LlamaIndex won't recognize → placeholder
        out = parse_to_text(b"\x00\x01\x02", "thing.xyz123")
        # Either placeholder or empty (LlamaIndex behavior on unknown ext)
        assert out == "" or out.startswith("[binary content")
        if out:
            assert "thing.xyz123" in out

    def test_no_extension(self):
        out = parse_to_text(b"\x00\x01\x02", "file_without_ext")
        assert out == "" or out.startswith("[binary content")

    def test_empty_input_returns_empty(self):
        # Empty bytes → empty string (regardless of LlamaIndex availability)
        assert parse_to_text(b"", "x.txt") == ""

    def test_empty_filename_no_crash(self):
        out = parse_to_text(b"hello", "")
        # No extension → unknown, but should not crash
        assert isinstance(out, str)
